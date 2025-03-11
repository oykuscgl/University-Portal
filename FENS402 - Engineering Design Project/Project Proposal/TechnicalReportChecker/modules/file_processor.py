from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os

def validate_report(filepath):
    doc = Document(filepath)
    corrected_doc = Document()
    error_report = []

    # JSON format kurallarını yükle
    with open('format_rules.json', 'r') as f:
        format_rules = json.load(f)

    # Paragrafları işlerken yeni bir dokümana aktarma
    for para in doc.paragraphs:
        new_para = corrected_doc.add_paragraph()  # Yeni paragraf oluştur

        # Başlık kontrolü ve ortalama
        if para.style.name.startswith('Heading'):
            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Başlıkları ortala
            run = new_para.add_run(para.text)
            run.bold = True  # Başlıkları kalın yap
            continue  # Başlıkları değiştirme, sadece ortalı ve kalın yap

        # Normal paragraf kontrolü
        for run in para.runs:
            new_run = new_para.add_run(run.text)  # Aynı metni yeni paragrafa ekle
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline

            # Font kontrolü
            if run.font.name and run.font.name != format_rules['font']['default']:
                error_report.append(f"Font hatası: '{run.text}' yanlış font kullanıyor.")
                new_run.font.color.rgb = RGBColor(255, 0, 0)  # Hatalı kelimeyi kırmızı yap

    # Düzeltilmiş dosyayı kaydet
    corrected_filename = "corrected_" + os.path.basename(filepath)
    corrected_filepath = os.path.join('uploads', corrected_filename)
    corrected_doc.save(corrected_filepath)

    # Hata raporunu kaydet
    error_report_filename = "error_report.txt"
    error_report_path = os.path.join('uploads', error_report_filename)
    with open(error_report_path, 'w') as f:
        for error in error_report:
            f.write(error + "\n")

    return corrected_filename, error_report_filename
